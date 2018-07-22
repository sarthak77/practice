"""
To add a line break (rather than starting a whole new paragraph), you
can call the add_break() method on the Run object you want to have
the break appear after. If you want to add a page break instead, you
need to pass the value docx.text.WD_BREAK.PAGE as a lone argument to
add_break(), as is done in the middle of the following example:
"""
import docx
doc = docx.Document()
doc.add_paragraph('Thi is on the first page!')
#error
doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')