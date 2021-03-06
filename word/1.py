"""
When using pip to first install Python-Docx, be sure to install
python-docx, not docx. The installation name docx is for a different
module that this book does not cover. However, when you are going to
import the python-docx module, you’ll need to run import docx, not
import python-docx.
"""
from docx import Document
doc = Document('demo.docx')
print(len(doc.paragraphs))
#   7
print(doc.paragraphs[0].text)
#   'Document Title'
print(doc.paragraphs[1].text)
#  'A plain paragraph with some bold and some italic'
print(len(doc.paragraphs[1].runs))
#   4
print(doc.paragraphs[1].runs[0].text)
#'A plain paragraph with some '
print(doc.paragraphs[1].runs[1].text)
#'bold'
print(doc.paragraphs[1].runs[2].text)
#' and some '
print(doc.paragraphs[1].runs[3].text)
#'italic'