"""
The arguments to add_heading() are a string of the heading text and
an integer from 0 to 4. The integer 0 makes the heading the Title
style, which is used for the top of the document. Integers 1 to 4 are
for various heading levels, with 1 being the main heading and 4 the
lowest subheading. The add_heading() function returns a Paragraph
object to save you the step of extracting it from the Document object
as a separate step.
"""

import docx
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')