"""
For Word documents, there are three types of styles: Paragraph styles
can be applied to Paragraph objects, character styles can be applied
to Run objects, and linked styles can be applied to both kinds of
objects. You can give both Paragraph and Run objects styles by
setting their style attribute to a string. This string should be the
name of a style. If style is set to None, then there will be no style
associated with the Paragraph or Run object.
"""
import docx
doc = docx.Document('demo.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text)
print((doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text))
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')